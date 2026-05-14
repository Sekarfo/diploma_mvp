"""
Vacancy file parser: extracts text from PDF / DOCX and infers structured fields.

Supported formats:
  • PDF  — via pdfplumber (text-based) with pypdf fallback
  • DOCX — via python-docx

Field extraction heuristics (no external AI required):
  • title    — first short line or line after known heading keywords
  • years    — regex over common patterns ("3+ years", "5 лет опыта", …)
  • skills   — skills-section parser + known-skills vocabulary scan
  • description — full extracted text (user can trim before submitting)
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import ClassVar


# ── Known tech skills vocabulary (case-insensitive whole-word match) ──────────

_KNOWN_SKILLS: frozenset[str] = frozenset({
    # Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "golang",
    "rust", "ruby", "scala", "kotlin", "swift", "php", "r", "matlab", "bash",
    "shell", "perl", "haskell", "lua", "dart", "elixir", "clojure",
    # Web / backend frameworks
    "django", "fastapi", "flask", "spring", "rails", "laravel", "express",
    "nestjs", "nuxt", "next.js", "nextjs", "react", "vue", "angular", "svelte",
    "htmx", "graphql", "rest", "grpc",
    # Data / ML / AI
    "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn", "xgboost",
    "lightgbm", "catboost", "pandas", "numpy", "scipy", "matplotlib", "seaborn",
    "spark", "pyspark", "hadoop", "hive", "kafka", "airflow", "dbt", "mlflow",
    "hugging face", "transformers", "langchain", "openai", "llm",
    # Databases
    "postgresql", "postgres", "mysql", "mariadb", "mongodb", "redis",
    "elasticsearch", "cassandra", "sqlite", "oracle", "mssql", "sql server",
    "clickhouse", "bigquery", "snowflake", "dynamodb",
    # Cloud / DevOps / Infra
    "aws", "azure", "gcp", "docker", "kubernetes", "k8s", "terraform",
    "ansible", "jenkins", "github actions", "gitlab ci", "ci/cd", "linux",
    "git", "nginx", "apache", "rabbitmq", "celery",
    # General engineering
    "microservices", "api", "oauth", "jwt", "agile", "scrum", "kanban",
    "tdd", "bdd", "solid", "oop", "design patterns",
})

# ── Regex patterns ─────────────────────────────────────────────────────────────

# Matches: "3+ years", "2-5 years", "5 лет", "minimum 4 years of experience"
_YEARS_RE = re.compile(
    r"(\d+)\s*(?:\+|\.0)?\s*(?:–|-|to|до)?\s*(?:\d+)?\s*"
    r"(?:\+\s*)?(?:years?|yr\.?s?|лет|года?)\s*"
    r"(?:of\s+)?(?:experience|опыта|exp\.?)?",
    re.IGNORECASE,
)

# Section headers that typically contain skill lists
_SKILLS_HEADER_RE = re.compile(
    r"^(?:"
    r"(?:required\s+)?skills?|requirements?|qualifications?|must.?have|"
    r"technical\s+skills?|technologies?|tech(?:nology)?\s+stack|stack|"
    r"what\s+you(?:['']ll)?\s+(?:need|bring|have|know)|"
    r"ключевые\s+навыки|требования|навыки|технологии"
    r")\s*:?\s*$",
    re.IGNORECASE,
)

# Line separators / bullets
_BULLET_RE = re.compile(r"^[\-•*·▪▸►✓✔>→]\s+")
_COMMA_LIKE_RE = re.compile(r"[,;/]")

# Title keywords that precede the actual title on the same line
_TITLE_PREFIX_RE = re.compile(
    r"^(?:job\s+title|position|vacancy|role|opening|job\s+posting|"
    r"вакансия|должность|позиция)\s*:?\s*",
    re.IGNORECASE,
)

# Lines that are almost certainly not a job title (too long or too generic)
_NOT_TITLE_RE = re.compile(
    r"^(?:about|introduction|overview|description|responsibilities|duties|"
    r"benefits|salary|compensation|location|remote|office|we\s+are|our\s+)",
    re.IGNORECASE,
)


@dataclass
class ParsedVacancy:
    title: str = ""
    description: str = ""
    years_required: float | None = None
    skills: list[str] = field(default_factory=list)
    file_name: str = ""
    char_count: int = 0
    page_count: int = 0
    parse_warnings: list[str] = field(default_factory=list)


class VacancyParserService:
    MAX_FILE_BYTES: ClassVar[int] = 10 * 1024 * 1024   # 10 MB
    ALLOWED_MIME_TYPES: ClassVar[frozenset[str]] = frozenset({
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    })
    ALLOWED_EXTENSIONS: ClassVar[frozenset[str]] = frozenset({".pdf", ".docx", ".doc"})

    # ── Public API ────────────────────────────────────────────────────────────

    def parse(self, content: bytes, file_name: str, content_type: str = "") -> ParsedVacancy:
        ext = self._ext(file_name)
        warnings: list[str] = []

        if len(content) > self.MAX_FILE_BYTES:
            raise ValueError(
                f"File too large ({len(content) // 1024} KB). Maximum is "
                f"{self.MAX_FILE_BYTES // 1024 // 1024} MB."
            )
        if ext not in self.ALLOWED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{ext}'. Allowed: "
                + ", ".join(sorted(self.ALLOWED_EXTENSIONS))
            )

        if ext == ".pdf":
            text, pages = self._extract_pdf(content, warnings)
        else:
            text, pages = self._extract_docx(content, warnings)

        if not text.strip():
            raise ValueError("Could not extract any text from the file. The file may be empty or image-only.")

        result = ParsedVacancy(
            file_name=file_name,
            char_count=len(text),
            page_count=pages,
            parse_warnings=warnings,
        )
        result.description = text
        result.title = self._extract_title(text)
        result.years_required = self._extract_years(text)
        result.skills = self._extract_skills(text)
        return result

    # ── Text extraction ───────────────────────────────────────────────────────

    def _extract_pdf(self, content: bytes, warnings: list[str]) -> tuple[str, int]:
        try:
            import pdfplumber
        except ImportError:
            warnings.append("pdfplumber not installed — falling back to pypdf.")
            return self._extract_pdf_fallback(content, warnings)

        try:
            pages_text: list[str] = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    t = page.extract_text(x_tolerance=2, y_tolerance=3)
                    if t:
                        pages_text.append(t)
            text = "\n".join(pages_text).strip()
            if not text:
                warnings.append("pdfplumber found no text — document may be scanned/image-based.")
            return text, page_count
        except Exception as exc:
            warnings.append(f"pdfplumber failed ({exc}), trying fallback.")
            return self._extract_pdf_fallback(content, warnings)

    def _extract_pdf_fallback(self, content: bytes, warnings: list[str]) -> tuple[str, int]:
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            pages_text = [p.extract_text() or "" for p in reader.pages]
            return "\n".join(pages_text).strip(), len(reader.pages)
        except ImportError:
            raise ImportError(
                "Neither pdfplumber nor pypdf is installed. "
                "Run: pip install pdfplumber"
            )

    def _extract_docx(self, content: bytes, warnings: list[str]) -> tuple[str, int]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is not installed. Run: pip install python-docx"
            )
        try:
            doc = Document(io.BytesIO(content))
            lines: list[str] = []
            prev_empty = False
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    lines.append(t)
                    prev_empty = False
                elif not prev_empty:
                    # preserve ONE empty line so section boundaries are detectable
                    lines.append("")
                    prev_empty = True
            # also pick up table cells
            seen = set(lines)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t and t not in seen:
                            lines.append(t)
                            seen.add(t)
            return "\n".join(lines), 1
        except Exception as exc:
            raise ValueError(f"Failed to parse DOCX: {exc}") from exc

    # ── Field extraction ──────────────────────────────────────────────────────

    def _extract_title(self, text: str) -> str:
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        # Pass 1: look for explicit heading keyword on a line
        for line in lines[:20]:
            m = _TITLE_PREFIX_RE.match(line)
            if m:
                after = line[m.end():].strip()
                if after and len(after) <= 120:
                    return after

        # Pass 2: first reasonably short line that doesn't look like boilerplate
        for line in lines[:10]:
            if _NOT_TITLE_RE.match(line):
                continue
            if _BULLET_RE.match(line):
                continue
            if len(line) <= 100 and len(line) >= 4:
                return line

        return lines[0][:100] if lines else ""

    def _extract_years(self, text: str) -> float | None:
        matches = _YEARS_RE.findall(text)
        if not matches:
            return None
        # Take the minimum numeric value found (entry-level requirement)
        values = [int(m) for m in matches if m.isdigit()]
        return float(min(values)) if values else None

    def _extract_skills(self, text: str) -> list[str]:
        found: set[str] = set()

        # Pass 1: look for explicit skills section and parse items from it
        lines = text.splitlines()
        in_skills_section = False
        for line in lines:
            stripped = line.strip()
            if _SKILLS_HEADER_RE.match(stripped):
                in_skills_section = True
                continue
            if in_skills_section:
                # leave section on empty line, known section header,
                # or any non-bullet line ending with ":" (e.g. "Nice to have:")
                if not stripped:
                    in_skills_section = False
                    continue
                if _SKILLS_HEADER_RE.match(stripped):
                    in_skills_section = False
                    continue
                if stripped.endswith(":") and not _BULLET_RE.match(stripped):
                    in_skills_section = False
                    continue
                # parse bullet items or comma-separated values
                item = _BULLET_RE.sub("", stripped)
                if _COMMA_LIKE_RE.search(item):
                    parts = _COMMA_LIKE_RE.split(item)
                    for p in parts:
                        s = p.strip().lower()
                        if s:
                            found.add(s)
                else:
                    s = item.strip().lower()
                    if s:
                        found.add(s)

        # Pass 2: scan full text for known skills vocabulary
        text_lower = text.lower()
        for skill in _KNOWN_SKILLS:
            pattern = rf"\b{re.escape(skill)}\b"
            if re.search(pattern, text_lower):
                found.add(skill)

        # Sub-split on "or" / "and" connectors that appear in bullet items
        expanded: set[str] = set()
        for item in found:
            if re.search(r"\b(?:or|and|либо|и)\b", item):
                for part in re.split(r"\b(?:or|and|либо|и)\b", item):
                    p = part.strip().strip("+-/.,;").strip()
                    if p:
                        expanded.add(p)
            else:
                expanded.add(item)

        # Keep only items that look like skill tokens (not sentences, not headers)
        _NOISE_RE = re.compile(
            r"(?:\d+\+?\s*years?|experience|requirements?|skills?:|"
            r"nice\s+to\s+have|responsibilities|benefits|\bwith\b|\busing\b)",
            re.IGNORECASE,
        )
        cleaned = sorted(
            s for s in expanded
            if 1 < len(s) <= 40
            and not s.startswith("http")
            and not s.endswith(":")
            and not s.startswith("(")
            and not s.endswith(")")             # trailing unmatched ")"
            and s.count("(") == s.count(")")   # balanced parentheses
            and not re.search(r"\s{2,}", s)    # no multi-space (headers)
            and not _NOISE_RE.search(s)
            and len(s.split()) <= 4            # max 4 words
        )
        return cleaned[:40]   # cap at 40 items

    @staticmethod
    def _ext(file_name: str) -> str:
        return ("." + file_name.rsplit(".", 1)[-1]).lower() if "." in file_name else ""
